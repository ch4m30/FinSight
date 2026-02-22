"""
Word document report generation using python-docx.
Produces a polished .docx working paper matching the PDF structure.

Structure:
  - Document header: "INTERNAL USE ONLY" on every page
  - Cover page: firm/client details + internal use notice
  - Executive Summary: health summary, snapshot metrics, red flags
  - Detailed sections: Profitability, Liquidity, Efficiency, Leverage, Growth
  - ATO Benchmark Comparison
  - AI Commentary (if provided)
  - Accountant's Notes (blank final page for handwritten notes)
"""

import io
import logging
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from utils.formatters import format_currency, format_percent, format_ratio, format_days

logger = logging.getLogger(__name__)

# ── Colours ──────────────────────────────────────────────────────────────────
NAVY_RGB  = RGBColor(0x1B, 0x2A, 0x4A)
GREY_RGB  = RGBColor(0x6B, 0x72, 0x80)
GREEN_RGB = RGBColor(0x16, 0xA3, 0x4A)
AMBER_RGB = RGBColor(0xD9, 0x77, 0x06)
RED_RGB   = RGBColor(0xDC, 0x26, 0x26)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)

STATUS_RGB = {
    "green": GREEN_RGB,
    "amber": AMBER_RGB,
    "red":   RED_RGB,
    "grey":  GREY_RGB,
}
STATUS_TEXT = {
    "green": "Good",
    "amber": "Review",
    "red":   "Concern",
    "grey":  "N/A",
}


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_para_border_bottom(para, color="1B2A4A", size="6"):
    """Add a bottom border to a paragraph (used for heading underlines)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_document_header(doc: Document, text: str):
    """Set the document header text on all pages."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        # Clear existing paragraphs
        for p in header.paragraphs:
            p.clear()
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RED_RGB


# ── Formatting helpers ────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading with navy colour and bottom border."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY_RGB
    if level <= 2:
        _set_para_border_bottom(p)


def _add_info_table(doc: Document, rows: list) -> None:
    """Add a 2-column label/value info table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for r_idx, (label, value) in enumerate(rows):
        cells = table.rows[r_idx].cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].width = Cm(5)
        cells[1].width = Cm(11)
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = NAVY_RGB
            run.font.size = Pt(10)
        for run in cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
        _set_cell_bg(cells[0], "F3F4F6")


def _add_metric_table(doc: Document, metrics: dict, category: str, period_labels: list) -> None:
    """Add a styled metric table for one category."""
    label_cur = period_labels[0] if period_labels else "Current"
    label_pri = period_labels[1] if len(period_labels) > 1 else "Prior"

    cat_metrics = [(k, v) for k, v in metrics.items() if v.category == category]
    if not cat_metrics:
        return

    table = doc.add_table(rows=1 + len(cat_metrics), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths
    for i, width in enumerate([Cm(7), Cm(3), Cm(3), Cm(2.5), Cm(2.5)]):
        for cell in table.columns[i].cells:
            cell.width = width

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_texts  = ["Metric", label_cur, label_pri, "Trend", "Status"]
    for cell, text in zip(hdr_cells, hdr_texts):
        cell.text = ""
        _set_cell_bg(cell, "1B2A4A")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = WHITE_RGB
        run.font.size = Pt(9)

    # Data rows
    status_fill = {
        "green": "DCFCE7",
        "amber": "FEF3C7",
        "red":   "FEE2E2",
        "grey":  "F3F4F6",
    }
    for row_idx, (key, m) in enumerate(cat_metrics, start=1):
        cells = table.rows[row_idx].cells
        data  = [m.label, m.current_fmt, m.prior_fmt, m.trend, STATUS_TEXT.get(m.status, "N/A")]
        for c_idx, (cell, text) in enumerate(zip(cells, data)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
            run.font.size = Pt(9)
            if c_idx == 4:
                _set_cell_bg(cell, status_fill.get(m.status, "F3F4F6"))
                run.font.color.rgb = STATUS_RGB.get(m.status, GREY_RGB)
                run.font.bold = True
        # Alternate row shading
        if row_idx % 2 == 0:
            for cell in cells[:4]:
                _set_cell_bg(cell, "F9FAFB")


def _add_bm_table(doc: Document, benchmark_comparisons: dict) -> None:
    """Add the ATO benchmark comparison table."""
    headers = ["Expense Category", "Client %", "ATO Low", "ATO High", "Status"]
    n_rows  = 1 + len(benchmark_comparisons)
    table   = doc.add_table(rows=n_rows, cols=5)
    table.style = "Table Grid"

    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = ""
        _set_cell_bg(cell, "1B2A4A")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = WHITE_RGB
        run.font.size = Pt(9)

    for r_idx, comp in enumerate(benchmark_comparisons.values(), start=1):
        actual   = comp.get("actual_pct")
        low      = comp.get("benchmark_low")
        high     = comp.get("benchmark_high")
        in_range = actual is not None and low is not None and high is not None and low <= actual <= high
        bm_fill  = "DCFCE7" if in_range else "FEF3C7"
        cells    = table.rows[r_idx].cells

        data = [
            comp["label"],
            f"{actual:.1f}%" if actual is not None else "N/A",
            f"{low:.0f}%"    if low    is not None else "N/A",
            f"{high:.0f}%"   if high   is not None else "N/A",
            "Within range" if in_range else "Outside range",
        ]
        for c_idx, (cell, text) in enumerate(zip(cells, data)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
            run.font.size = Pt(9)
        _set_cell_bg(cells[4], bm_fill)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_word_report(
    analysis_result,
    session_info: dict,
    commentary: str,
    firm_name: str = "Your Firm Name",
) -> bytes:
    """
    Generate a polished Word document and return as bytes.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Persistent header
    _add_document_header(doc, "INTERNAL USE ONLY — NOT FOR DISTRIBUTION")

    client_name   = session_info.get("client_name", "Client")
    abn           = session_info.get("abn", "")
    industry      = session_info.get("industry", "")
    fy_end        = session_info.get("financial_year_end", "")
    currency      = session_info.get("currency", "AUD")
    analysis_date = date.today().strftime("%d %B %Y")
    period_labels = analysis_result.period_labels
    label_cur     = period_labels[0] if period_labels else "Current"
    label_pri     = period_labels[1] if len(period_labels) > 1 else "Prior"

    # ── COVER PAGE ───────────────────────────────────────────────────────────

    # Firm name
    firm_para = doc.add_paragraph()
    firm_run  = firm_para.add_run(firm_name)
    firm_run.font.size  = Pt(22)
    firm_run.font.bold  = True
    firm_run.font.color.rgb = NAVY_RGB

    # Report title
    title_para = doc.add_paragraph()
    title_run  = title_para.add_run("Financial Analysis Report")
    title_run.font.size  = Pt(14)
    title_run.font.color.rgb = GREY_RGB

    doc.add_paragraph()

    # Client info table
    _add_info_table(doc, [
        ("Client / Business Name", client_name),
        ("ABN",                    abn or "Not provided"),
        ("Industry",               industry),
        ("Financial Year End",     fy_end),
        ("Reporting Currency",     currency),
        ("Date of Analysis",       analysis_date),
        ("Current Period",         label_cur),
        ("Prior Period",           label_pri),
        ("Prepared by",            "FinSight — Internal Use Only"),
    ])

    doc.add_paragraph()

    # Internal use notice
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice.add_run(
        "INTERNAL USE ONLY — This document is prepared for internal accountant use. "
        "Do not distribute without appropriate review and professional sign-off."
    )
    notice_run.font.bold  = True
    notice_run.font.color.rgb = RED_RGB
    notice_run.font.size  = Pt(9)

    # Track changes note
    tc_note = doc.add_paragraph()
    tc_run  = tc_note.add_run(
        "Note: This document is Track Changes ready. Enable Track Changes before editing "
        "to maintain an audit trail of modifications."
    )
    tc_run.font.size  = Pt(8)
    tc_run.font.color.rgb = GREY_RGB
    tc_run.font.italic    = True

    doc.add_page_break()

    # ── EXECUTIVE SUMMARY ────────────────────────────────────────────────────
    _add_heading(doc, "Executive Summary", level=1)

    # Health counts
    g_count = sum(1 for m in analysis_result.metrics.values() if m.status == "green")
    a_count = sum(1 for m in analysis_result.metrics.values() if m.status == "amber")
    r_count = sum(1 for m in analysis_result.metrics.values() if m.status == "red")

    health_tbl = doc.add_table(rows=1, cols=4)
    health_tbl.style = "Table Grid"
    health_data = [
        ("Overall Health", "1B2A4A", WHITE_RGB),
        (f"{g_count} Good",    "DCFCE7", GREEN_RGB),
        (f"{a_count} Review",  "FEF3C7", AMBER_RGB),
        (f"{r_count} Concern", "FEE2E2", RED_RGB),
    ]
    for i, (text, bg, fg) in enumerate(health_data):
        cell = health_tbl.rows[0].cells[i]
        cell.text = ""
        _set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = fg
        run.font.size = Pt(10)

    doc.add_paragraph()
    _add_heading(doc, "Key Metrics Snapshot", level=2)

    # Snapshot: build table directly across all categories
    spotlight_keys = [
        "gross_profit_margin", "net_profit_margin", "ebitda_margin",
        "current_ratio", "debtor_days", "debt_to_equity",
    ]
    status_fill = {"green": "DCFCE7", "amber": "FEF3C7", "red": "FEE2E2", "grey": "F3F4F6"}

    snap_table = doc.add_table(rows=1, cols=5)
    snap_table.style = "Table Grid"
    for cell, text in zip(snap_table.rows[0].cells,
                          ["Metric", label_cur, label_pri, "Trend", "Status"]):
        cell.text = ""
        _set_cell_bg(cell, "1B2A4A")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.bold = True
        run.font.color.rgb = WHITE_RGB
        run.font.size = Pt(9)

    for key in spotlight_keys:
        m = analysis_result.metrics.get(key)
        if m is None:
            continue
        row = snap_table.add_row()
        data = [m.label, m.current_fmt, m.prior_fmt, m.trend, STATUS_TEXT.get(m.status, "N/A")]
        for c_idx, (cell, text) in enumerate(zip(row.cells, data)):
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
            run.font.size = Pt(9)
            if c_idx == 4:
                _set_cell_bg(cell, status_fill.get(m.status, "F3F4F6"))
                run.font.color.rgb = STATUS_RGB.get(m.status, GREY_RGB)
                run.font.bold = True

    doc.add_paragraph()
    _add_heading(doc, "Red Flags", level=2)
    if analysis_result.red_flags:
        for flag in analysis_result.red_flags:
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ {flag}")
            run.font.color.rgb = RED_RGB
            run.font.size = Pt(10)
    else:
        p = doc.add_paragraph("No red flags detected.")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── DETAILED SECTIONS ────────────────────────────────────────────────────
    categories = [
        ("profitability", "Profitability"),
        ("liquidity",     "Liquidity & Working Capital"),
        ("efficiency",    "Operational Efficiency"),
        ("leverage",      "Leverage & Solvency"),
        ("growth",        "Growth Metrics"),
    ]

    for cat_key, cat_label in categories:
        cat_metrics = {k: v for k, v in analysis_result.metrics.items() if v.category == cat_key}
        if not cat_metrics:
            continue

        _add_heading(doc, cat_label, level=1)
        _add_metric_table(doc, analysis_result.metrics, cat_key, period_labels)

        for m in cat_metrics.values():
            if m.notes:
                p = doc.add_paragraph()
                run = p.add_run(f"⚠ {m.notes}")
                run.font.color.rgb = RED_RGB
                run.font.size = Pt(9)

        doc.add_paragraph()

    # ── ATO BENCHMARKS ───────────────────────────────────────────────────────
    if analysis_result.benchmark_comparisons:
        doc.add_page_break()
        _add_heading(doc, "ATO Small Business Benchmark Comparison", level=1)

        note = doc.add_paragraph(
            f"Industry: {industry}. Benchmarks expressed as % of turnover. "
            "Note: ATO benchmarks are updated annually and may lag by one financial year."
        )
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = GREY_RGB

        _add_bm_table(doc, analysis_result.benchmark_comparisons)
        doc.add_paragraph()

        src = doc.add_paragraph(
            "Source: ATO Small Business Benchmarks (ato.gov.au)"
        )
        src.runs[0].font.size = Pt(8)
        src.runs[0].font.color.rgb = GREY_RGB
        src.runs[0].font.italic = True

    # ── AI COMMENTARY ────────────────────────────────────────────────────────
    if commentary:
        doc.add_page_break()
        _add_heading(doc, "Financial Commentary", level=1)

        disclaimer = doc.add_paragraph(
            "The following commentary has been generated by a local AI model to assist in "
            "client meeting preparation. Review and edit as required before client delivery. "
            "This commentary does not constitute professional financial advice."
        )
        disclaimer.runs[0].font.size = Pt(9)
        disclaimer.runs[0].font.color.rgb = GREY_RGB
        disclaimer.runs[0].font.italic = True
        doc.add_paragraph()

        for line in commentary.split("\n"):
            line_s = line.strip()
            if not line_s:
                doc.add_paragraph()
            elif line_s.startswith("## "):
                _add_heading(doc, line_s[3:], level=2)
            elif line_s.startswith("### "):
                _add_heading(doc, line_s[4:], level=3)
            elif line_s.startswith("- ") or line_s.startswith("* "):
                p = doc.add_paragraph(line_s[2:], style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(10)
            else:
                p = doc.add_paragraph(line_s)
                for run in p.runs:
                    run.font.size = Pt(10)

    # ── ACCOUNTANT'S NOTES ───────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "Accountant's Notes", level=1)

    notes_intro = doc.add_paragraph(
        "Use this page for handwritten or typed notes during client review. "
        "Additional observations, follow-up items, and action points:"
    )
    notes_intro.runs[0].font.size = Pt(10)
    notes_intro.runs[0].font.color.rgb = GREY_RGB

    # Add lined space for notes
    for _ in range(20):
        rule_p = doc.add_paragraph()
        rule_p.paragraph_format.space_after = Pt(16)
        _set_para_border_bottom(rule_p, color="D1D5DB", size="4")

    # Document footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Prepared by FinSight | {analysis_date} | Internal Use Only"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = GREY_RGB
    footer_p.runs[0].font.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
